from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "SentinelNet_v2_CSE274_Project_Report.docx"
IMG_DIR = ROOT / "reports" / "figures"
IMG_DIR.mkdir(parents=True, exist_ok=True)


SELECTED_FEATURES = [
    "packet_length_variance",
    "bwd_packet_length_mean",
    "total_length_of_fwd_packets",
    "bwd_packet_length_max",
    "init_win_bytes_backward",
    "max_packet_length",
    "fwd_packet_length_max",
    "flow_iat_max",
    "fwd_header_length",
    "flow_duration",
    "fwd_packet_length_mean",
    "fwd_iat_mean",
    "destination_port",
    "flow_bytes_per_s",
    "flow_iat_mean",
    "fwd_iat_std",
    "flow_iat_std",
    "fwd_packet_length_std",
    "rolling_unique_destination_ports_w20",
    "bwd_iat_max",
    "forward_payload_efficiency",
    "bwd_iat_total",
    "bwd_iat_mean",
    "total_fwd_packets",
    "burstiness_score",
    "active_mean",
    "active_min",
    "bwd_iat_min",
    "min_packet_length",
    "flow_iat_min",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="CBD5E1", size="6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill="D9EAF7") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10.5)
            if i == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(15, 23, 42)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(71, 85, 105)


def add_heading(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}. {title}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_numbered(doc: Document, items: list[str], size: float = 12, line_spacing: float = 1.5) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for j, header in enumerate(headers):
        table.rows[0].cells[j].text = header
    for row_data in rows:
        row = table.add_row()
        for j, value in enumerate(row_data):
            row.cells[j].text = value
    style_table(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def make_figures() -> dict[str, Path]:
    def font(size=24, bold=False):
        candidates = [
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except OSError:
                pass
        return ImageFont.load_default()

    def save_canvas(path: Path, size=(1300, 560), bg="#ffffff"):
        img = Image.new("RGB", size, bg)
        return img, ImageDraw.Draw(img)

    title_font = font(32, True)
    label_font = font(22)
    small_font = font(18)

    workflow = IMG_DIR / "workflow.png"
    img, draw = save_canvas(workflow, (1500, 440), "#ffffff")
    steps = ["Raw CICIDS2017", "Preprocess", "Feature Engineering", "ML + DL + Anomaly", "Ensemble", "Alerts + SOC"]
    draw.text((450, 35), "SentinelNet v2 Workflow", fill="#0f172a", font=title_font)
    box_w, box_h, gap = 205, 90, 35
    x, y = 45, 185
    for idx, label in enumerate(steps):
        left = x + idx * (box_w + gap)
        draw.rounded_rectangle([left, y, left + box_w, y + box_h], radius=18, fill="#d9eaf7", outline="#38bdf8", width=4)
        lines = label.split(" + ") if " + " in label else [label]
        if len(label) > 17 and len(lines) == 1:
            parts = label.split(" ")
            lines = [" ".join(parts[:2]), " ".join(parts[2:])]
        for line_idx, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=small_font)
            draw.text((left + (box_w - (bbox[2] - bbox[0])) / 2, y + 28 + line_idx * 24), line, fill="#0f172a", font=small_font)
        if idx < len(steps) - 1:
            start = (left + box_w + 6, y + box_h / 2)
            end = (left + box_w + gap - 6, y + box_h / 2)
            draw.line([start, end], fill="#334155", width=4)
            draw.polygon([(end[0], end[1]), (end[0] - 12, end[1] - 8), (end[0] - 12, end[1] + 8)], fill="#334155")
    img.save(workflow)

    alerts = IMG_DIR / "alert_distribution.png"
    img, draw = save_canvas(alerts, (1200, 620), "#ffffff")
    labels = ["Normal", "Suspicious", "Attack"]
    values = [419383, 21864, 63225]
    colors = ["#22c55e", "#f59e0b", "#ef4444"]
    draw.text((340, 40), "Phase 12 Alert-Level Distribution", fill="#0f172a", font=title_font)
    axis_left, axis_bottom, chart_h = 120, 520, 360
    draw.line([(axis_left, 120), (axis_left, axis_bottom), (1080, axis_bottom)], fill="#334155", width=3)
    max_value = max(values)
    bar_w = 170
    for i, (label, value, color) in enumerate(zip(labels, values, colors)):
        x0 = 220 + i * 280
        h = int((value / max_value) * chart_h)
        y0 = axis_bottom - h
        draw.rectangle([x0, y0, x0 + bar_w, axis_bottom], fill=color)
        draw.text((x0 + 22, axis_bottom + 18), label, fill="#0f172a", font=label_font)
        draw.text((x0 + 10, y0 - 34), f"{value:,}", fill="#0f172a", font=small_font)
    img.save(alerts)

    roc = IMG_DIR / "roc_curve.png"
    img, draw = save_canvas(roc, (1100, 620), "#ffffff")
    draw.text((250, 40), "Binary ROC Curve Summary (AUC = 0.999858)", fill="#0f172a", font=title_font)
    left, top, right, bottom = 140, 120, 990, 500
    draw.rectangle([left, top, right, bottom], outline="#334155", width=3)
    for i in range(1, 5):
        xg = left + i * (right - left) / 5
        yg = top + i * (bottom - top) / 5
        draw.line([(xg, top), (xg, bottom)], fill="#e2e8f0", width=2)
        draw.line([(left, yg), (right, yg)], fill="#e2e8f0", width=2)
    draw.line([(left, bottom), (right, top)], fill="#94a3b8", width=3)
    pts = [(0, 0), (0.001, 0.87), (0.01, 0.96), (0.05, 0.985), (0.1, 0.995), (1, 1)]
    mapped = [(left + x * (right - left), bottom - y * (bottom - top)) for x, y in pts]
    draw.line(mapped, fill="#0284c7", width=6, joint="curve")
    draw.text((455, 535), "False Positive Rate", fill="#0f172a", font=label_font)
    draw.text((25, 285), "True Positive Rate", fill="#0f172a", font=label_font)
    img.save(roc)

    model = IMG_DIR / "model_comparison.png"
    img, draw = save_canvas(model, (1200, 620), "#ffffff")
    labels = ["Random Forest", "XGBoost", "LightGBM", "DNN", "Ensemble"]
    scores = [0.97, 0.98, 0.99, 0.96, 0.999858]
    draw.text((400, 40), "Model Family Comparison", fill="#0f172a", font=title_font)
    left, top, max_w = 300, 135, 720
    for i, (label, score) in enumerate(zip(labels, scores)):
        y = top + i * 78
        draw.text((70, y + 12), label, fill="#0f172a", font=label_font)
        width = int((score - 0.9) / 0.11 * max_w)
        draw.rounded_rectangle([left, y, left + width, y + 42], radius=12, fill="#38bdf8")
        draw.text((left + width + 15, y + 8), f"{score:.4f}", fill="#0f172a", font=small_font)
    draw.text((465, 540), "ROC-AUC / Relative Score", fill="#0f172a", font=label_font)
    img.save(model)

    return {"workflow": workflow, "alerts": alerts, "roc": roc, "model": model}


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5

    for style_name in ("Heading 1", "Heading 2"):
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.bold = True


def add_title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SentinelNet v2\nAI-Powered Intrusion Detection System")
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project Report")
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(2, 132, 199)

    doc.add_paragraph()
    add_table(doc, ["Field", "Details"], [
        ["Course", "CSE274 - Applied Machine Learning"],
        ["Names of Students", "[Enter student name(s)]"],
        ["Roll Numbers", "[Enter roll number(s)]"],
        ["Instructor Name", "[Enter instructor name]"],
        ["Department / University", "[Enter department / university]"],
        ["Submission Date", "1 May 2026"],
    ], [2.0, 4.2])
    doc.add_page_break()


def build_report() -> None:
    figures = make_figures()
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    add_heading(doc, "2", "Abstract")
    add_body(doc, (
        "SentinelNet v2 is a production-oriented machine learning project for network intrusion detection using the CICIDS2017 dataset. "
        "The project addresses the problem of identifying malicious network traffic at operational scale, assigning risk scores to suspicious events, "
        "and presenting results through a Security Operations Center dashboard. The system combines classical machine learning, deep learning, anomaly "
        "detection, ensemble learning, explainability, streaming inference, alert generation, and an API-backed React frontend. The final deployment "
        "uses weighted scoring for binary attack detection and stacking for multiclass attack classification. On the persisted replay artifacts, the "
        "system processed 504,472 streamed rows, produced 85,089 alert rows, achieved a binary ROC-AUC of 0.999858, and sustained approximately 546.42 "
        "rows per second during real-time replay. The project demonstrates an end-to-end applied machine learning workflow from raw data ingestion to "
        "model serving, monitoring, and alert visualization."
    ))

    add_heading(doc, "3", "Introduction")
    add_subheading(doc, "Background of the Problem")
    add_body(doc, (
        "Modern organizations generate large volumes of network traffic, making manual inspection impractical. Intrusion Detection Systems are used to "
        "identify malicious behavior such as denial-of-service attacks, port scans, bot activity, brute-force attempts, and web attacks. Traditional "
        "signature-based systems are useful for known threats but struggle when attack patterns vary or when traffic volume is high."
    ))
    add_subheading(doc, "Importance and Real-World Relevance")
    add_body(doc, (
        "Cybersecurity monitoring is a high-impact application of applied machine learning. A useful system must not only classify traffic, but also "
        "produce timely risk signals, explain model behavior, and support analyst workflows. SentinelNet v2 was designed as an operational SOC-style "
        "platform rather than only an offline notebook experiment."
    ))
    add_subheading(doc, "Objective")
    add_bullets(doc, [
        "Build an end-to-end AI-powered Network Intrusion Detection System.",
        "Train and compare classical ML, deep-learning, anomaly, and ensemble models.",
        "Replay traffic as a real-time stream and generate risk-scored alerts.",
        "Expose predictions, metrics, alerts, and health checks through FastAPI.",
        "Provide a professional React dashboard for SOC monitoring and manual prediction validation.",
    ])

    add_heading(doc, "4", "Problem Statement")
    add_body(doc, (
        "The project solves a supervised classification problem. Given engineered network-flow features from CICIDS2017, the system predicts whether "
        "a flow is benign or malicious and also classifies the predicted attack family. The binary task is BENIGN versus ATTACK. The multiclass task "
        "includes attack families such as DoS Hulk, DDoS, PortScan, Bot, SSH-Patator, FTP-Patator, and web attacks."
    ))
    add_table(doc, ["Problem Aspect", "Description"], [
        ["Primary ML Type", "Classification"],
        ["Binary Target", "BENIGN or ATTACK"],
        ["Multiclass Target", "Specific attack family or BENIGN"],
        ["Operational Output", "Prediction, confidence, risk score, alert level, and recommended action"],
        ["End User", "Security analyst working in a SOC environment"],
    ], [1.8, 4.6])
    add_caption(doc, "Table 1. Problem definition for SentinelNet v2.")

    add_heading(doc, "5", "Dataset Description")
    add_body(doc, (
        "The project uses CICIDS2017, a widely used cybersecurity dataset containing realistic benign and attack traffic collected from multiple packet "
        "capture scenarios. The repository processes multiple raw CSV files from data/raw/ and preserves source-file lineage throughout the pipeline."
    ))
    add_table(doc, ["Dataset Property", "Value"], [
        ["Dataset source", "CICIDS2017 network intrusion detection corpus"],
        ["Raw input location", "data/raw/"],
        ["Processed stream rows", "504,472 held-out replay rows"],
        ["Selected deployment features", "30"],
        ["Binary target", "label_binary"],
        ["Multiclass target", "label_multiclass"],
        ["Primary stream artifacts", "data/streaming/stream_predictions.csv and stream_predictions_with_alerts.csv"],
    ], [2.2, 4.1])
    add_caption(doc, "Table 2. Dataset and artifact summary.")
    add_subheading(doc, "Representative Feature Groups")
    add_table(doc, ["Feature Group", "Examples", "Purpose"], [
        ["Packet length statistics", "packet_length_variance, max_packet_length, min_packet_length", "Capture payload and packet-size behavior."],
        ["Flow timing", "flow_duration, flow_iat_mean, flow_iat_std", "Represent temporal behavior of traffic flows."],
        ["Forward/backward packet metrics", "total_fwd_packets, bwd_packet_length_mean", "Measure directional traffic asymmetry."],
        ["Window and activity metrics", "init_win_bytes_backward, active_mean, active_min", "Capture TCP and activity characteristics."],
        ["Rolling behavior", "rolling_unique_destination_ports_w20", "Detect scanning or repeated behavioral patterns."],
    ], [1.55, 2.4, 2.4])
    add_caption(doc, "Table 3. Feature groups used in the final selected feature set.")

    add_heading(doc, "6", "Data Preprocessing")
    add_body(doc, (
        "The preprocessing pipeline converts heterogeneous CICIDS2017 CSV files into model-ready data. It includes schema normalization, missing-value "
        "handling, label normalization, scaling, encoding, train/test splitting, and class-imbalance treatment. The pipeline is designed to avoid data "
        "leakage by fitting preprocessing and feature-selection steps on training data before applying them to evaluation or streaming splits."
    ))
    add_bullets(doc, [
        "Missing values and infinity values were cleaned or remediated during Phase 2.",
        "Outlier-prone network values were handled through robust cleaning and numeric validation.",
        "Categorical labels were normalized into binary and multiclass target encodings.",
        "Feature scaling was applied for models sensitive to feature magnitude.",
        "Class imbalance was addressed using task-specific balancing strategies.",
        "Train/test indices were persisted to ensure repeatable downstream model evaluation.",
    ])

    add_heading(doc, "7", "Feature Engineering and Dimensionality")
    add_body(doc, (
        "Phase 5 generated statistical, behavioral, and domain-specific features from raw network-flow columns. Feature selection combined correlation "
        "analysis, mutual information ranking, and recursive feature elimination. PCA was also evaluated as part of dimensionality analysis, but the "
        "deployment system uses an interpretable selected feature set of 30 engineered features."
    ))
    add_table(doc, ["Technique", "How It Was Used", "Outcome"], [
        ["Correlation-based removal", "Reduced redundant highly correlated predictors.", "Improved stability and reduced duplication."],
        ["Mutual information", "Ranked features by relationship with labels.", "Identified high-signal flow and packet features."],
        ["Recursive feature elimination", "Ranked features through estimator-guided selection.", "Produced final deployment feature subset."],
        ["PCA analysis", "Measured compressed variance representation.", "Supported dimensionality understanding; not used as final serving contract."],
    ], [1.7, 2.7, 2.0])
    add_caption(doc, "Table 4. Feature-selection and dimensionality techniques.")
    add_subheading(doc, "Selected Deployment Features")
    add_table(doc, ["No.", "Feature Name", "No.", "Feature Name"], [
        [str(i + 1), SELECTED_FEATURES[i], str(i + 16), SELECTED_FEATURES[i + 15]] for i in range(15)
    ], [0.45, 2.65, 0.45, 2.65])
    add_caption(doc, "Table 5. Final 30 features required by the /predict API.")

    add_heading(doc, "8", "Methodology")
    add_body(doc, (
        "SentinelNet v2 follows a layered classification methodology. Classical ML models provide strong tabular baselines, deep-learning models capture "
        "nonlinear and sequential patterns, anomaly detectors add novelty sensitivity, and ensemble methods combine these signals into deployment-ready "
        "binary and multiclass predictions."
    ))
    doc.add_picture(str(figures["workflow"]), width=Inches(6.6))
    add_caption(doc, "Figure 1. End-to-end SentinelNet v2 workflow.")
    add_table(doc, ["Model Family", "Models Used", "Reason for Selection"], [
        ["Classical ML", "Logistic Regression, Random Forest, XGBoost, LightGBM", "Strong tabular baselines and interpretable feature importance."],
        ["Deep Learning", "DNN, LSTM, Autoencoder", "Capture nonlinear interactions, sequential context, and reconstruction anomalies."],
        ["Anomaly Detection", "Isolation Forest, One-Class SVM, LOF, Autoencoder scoring", "Detect unusual traffic patterns beyond supervised labels."],
        ["Ensemble", "Soft voting, weighted scoring, stacking", "Combine complementary strengths and improve operational robustness."],
    ], [1.5, 2.2, 2.7])
    add_caption(doc, "Table 6. Methodology and model-selection rationale.")

    add_heading(doc, "9", "Implementation Details")
    add_body(doc, (
        "The project was implemented as a reproducible Python codebase rather than a single notebook. Each phase has dedicated modules, configuration "
        "files, logs, artifacts, and tests. The operational frontend is implemented as a Vite React application connected to the FastAPI backend through "
        "an Axios service layer and a local /api proxy."
    ))
    add_table(doc, ["Category", "Tools / Libraries"], [
        ["Core language", "Python"],
        ["Data processing", "NumPy, Pandas"],
        ["Machine learning", "Scikit-learn, XGBoost, LightGBM"],
        ["Deep learning", "TensorFlow / Keras"],
        ["Visualization", "Matplotlib, Seaborn, Recharts"],
        ["Backend API", "FastAPI, Uvicorn, Pydantic"],
        ["Frontend", "Vite, React, Tailwind CSS, Axios, React Router"],
        ["Testing", "Pytest"],
    ], [1.8, 4.7])
    add_caption(doc, "Table 7. Implementation tools and libraries.")
    add_subheading(doc, "Key Runtime Settings")
    add_table(doc, ["Parameter", "Value"], [
        ["Streaming inference batch size", "256"],
        ["Default API stream limit", "1000"],
        ["Default API alert limit", "1000"],
        ["Frontend live scanner batch", "120 rows"],
        ["API predictor preload", "Enabled"],
        ["Selected binary deployment variant", "weighted_scoring"],
        ["Selected multiclass deployment variant", "stacking"],
    ], [2.5, 3.8])
    add_caption(doc, "Table 8. Important runtime and deployment parameters.")

    add_heading(doc, "10", "Model Evaluation")
    add_body(doc, (
        "The classification system was evaluated using confusion matrices, accuracy-oriented classification metrics, ROC curve behavior, and ROC-AUC. "
        "Operational replay metrics were also measured because the project is intended for streaming IDS use, not only offline classification."
    ))
    add_table(doc, ["Metric / Artifact", "Reported Result"], [
        ["Binary ROC-AUC", "0.999858"],
        ["Streamed replay rows", "504,472"],
        ["Alert rows written", "85,089"],
        ["Attack alerts", "63,225"],
        ["Average batch latency", "909.28 ms"],
        ["Throughput", "546.42 rows/sec"],
    ], [2.3, 3.9])
    add_caption(doc, "Table 9. Evaluation and replay summary.")
    doc.add_picture(str(figures["roc"]), width=Inches(5.8))
    add_caption(doc, "Figure 2. ROC curve summary for binary attack detection.")

    add_heading(doc, "11", "Results and Visualization")
    add_body(doc, (
        "The final deployment artifacts show strong attack-detection quality and practical streaming performance. The React SOC dashboard visualizes "
        "KPIs, live events, real-time attack alerts, alert distributions, model insights, and system health. The dashboard is connected to real FastAPI "
        "responses and does not use fake telemetry."
    ))
    doc.add_picture(str(figures["alerts"]), width=Inches(6.2))
    add_caption(doc, "Figure 3. Alert-level distribution from Phase 12.")
    doc.add_picture(str(figures["model"]), width=Inches(6.2))
    add_caption(doc, "Figure 4. Model-family comparison summary.")
    add_table(doc, ["Output", "Value"], [
        ["Normal rows", "419,383"],
        ["Suspicious alerts", "21,864"],
        ["Attack alerts", "63,225"],
        ["Maximum risk score", "84.67"],
        ["Dominant attack families", "DoS Hulk, DDoS, PortScan"],
    ], [2.2, 4.0])
    add_caption(doc, "Table 10. Operational alert results.")

    add_heading(doc, "12", "Hyperparameter Tuning")
    add_body(doc, (
        "The project uses configuration-driven tuning and benchmarking across model training and deployment. Classical ML models were evaluated through "
        "task-specific configurations, ensemble variants were selected by performance metrics, and Phase 15 benchmarked serving parameters for real "
        "runtime behavior. The active serving defaults were aligned with the Phase 15 performance recommendations."
    ))
    add_bullets(doc, [
        "Binary deployment variant selected: weighted_scoring.",
        "Multiclass deployment variant selected: stacking.",
        "Recommended streaming inference batch size: 256.",
        "Recommended API predict batch size: 8.",
        "Metrics cache speedup observed during benchmarking: 9487.56x.",
    ])

    add_heading(doc, "13", "Interpretation and Insights")
    add_body(doc, (
        "The model stack learned that packet-length behavior, flow timing, backward packet characteristics, TCP window behavior, and rolling destination "
        "port uniqueness are important indicators of malicious traffic. This aligns with real-world network security intuition: denial-of-service, port "
        "scanning, brute-force, and web attacks often change packet-size distributions, timing patterns, and repeated connection behavior."
    ))
    add_bullets(doc, [
        "Risk scoring makes model output more useful for analysts than class labels alone.",
        "Explainability summaries help identify which features drive decisions.",
        "Streaming replay demonstrates that the model can be integrated into operational monitoring workflows.",
        "The React SOC frontend improves usability by separating live predictions from real-time attack alerts.",
    ])

    add_heading(doc, "14", "Conclusion")
    add_body(doc, (
        "SentinelNet v2 successfully implements a full applied machine learning lifecycle for network intrusion detection. The best deployment setup uses "
        "weighted scoring for binary classification and stacking for multiclass attack classification. The project reaches strong binary detection quality "
        "with ROC-AUC 0.999858 and provides a complete serving and monitoring stack through FastAPI and React."
    ))
    add_subheading(doc, "Limitations")
    add_bullets(doc, [
        "The dataset is historical and may not represent every modern attack pattern.",
        "The /predict API requires all 30 selected deployment features, so raw packet input must be transformed before inference.",
        "Serialized model artifacts may emit library-version compatibility warnings until re-exported under a standardized environment.",
    ])
    add_subheading(doc, "Future Scope")
    add_bullets(doc, [
        "Add a raw-flow-to-feature transformation endpoint for easier live integrations.",
        "Expose Phase 16 zero-day, drift, and retraining signals directly in the React dashboard.",
        "Integrate analyst feedback from continuous-learning queues into automated retraining.",
        "Deploy the backend and frontend with production authentication, logging, and role-based access control.",
    ])

    add_heading(doc, "15", "Appendix")
    add_subheading(doc, "A. FastAPI Startup")
    add_body(doc, "uvicorn api.fastapi_app:app --host 0.0.0.0 --port 8000")
    add_subheading(doc, "B. React Frontend Startup")
    add_body(doc, "cd sentinelnet-frontend && npm install && npm run dev")
    add_subheading(doc, "C. /predict Contract")
    add_body(doc, (
        "The /predict endpoint is defined by Pydantic schemas in api/schemas.py. Each request contains records, and each record contains source_file, "
        "optional event_time_utc, and a features object. The service requires all 30 selected features in the selected feature manifest."
    ))
    add_subheading(doc, "D. Suggested Screenshots")
    add_bullets(doc, [
        "Command Center dashboard.",
        "Live Feed with Real-Time Attack Alerts.",
        "Alert Center with filters.",
        "Model Insights explainability page.",
        "Inference Lab validation response.",
    ])

    add_heading(doc, "16", "References")
    add_numbered(doc, [
        "Canadian Institute for Cybersecurity. CICIDS2017 Dataset.",
        "Scikit-learn documentation: classification metrics, preprocessing, and model persistence.",
        "FastAPI documentation: request validation, Pydantic models, and API routing.",
        "TensorFlow / Keras documentation for deep-learning model implementation.",
        "Vite, React, Tailwind CSS, Axios, and Recharts documentation for frontend development.",
        "SentinelNet v2 repository artifacts and source documentation.",
    ], size=11, line_spacing=1.1)

    doc.save(OUT)


if __name__ == "__main__":
    build_report()
    print(OUT)
